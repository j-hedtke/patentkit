"""Tests for chart/report rendering (DOCX test skips without python-docx)."""

import pytest

from patentkit.analysis.fto import FtoFinding, FtoReportData
from patentkit.analysis.infringement import EvidenceItem, InfringementFinding
from patentkit.analysis.invalidity import ClaimChart, DisclosureFinding, ReferenceChart
from patentkit.formatting import (
    claim_chart_html,
    claim_chart_markdown,
    fto_report_markdown,
    infringement_report_markdown,
    invalidity_report_markdown,
)
from patentkit.models import Limitation, Patent, PatentNumber
from patentkit.search.base import Passage, SearchResult


def make_chart() -> ClaimChart:
    lim_a = Limitation(label="1[a]", text="a frame supporting the housing")
    lim_b = Limitation(label="1[b]", text="a motor coupled to the frame")
    return ClaimChart(
        query_patent="US10123456B2",
        claim_number=1,
        limitations=[lim_a, lim_b],
        references=[
            ReferenceChart(
                reference_number="US111",
                reference_title="Prior Widget",
                findings=[
                    DisclosureFinding(
                        limitation=lim_a,
                        status="disclosed",
                        reasoning="The frame is taught.",
                        quotes=["a frame 12 supports the housing 14"],
                        citation="col. 2, ll. 10-12",
                    ),
                    DisclosureFinding(
                        limitation=lim_b, status="not_disclosed", reasoning="No motor."
                    ),
                ],
            ),
            ReferenceChart(
                reference_number="US222",
                findings=[
                    DisclosureFinding(limitation=lim_a, status="partial", reasoning="Implied."),
                    DisclosureFinding(
                        limitation=lim_b, status="disclosed", reasoning="Motor shown."
                    ),
                ],
            ),
        ],
    )


class TestClaimChartMarkdown:
    def test_contains_expected_cells(self):
        md = claim_chart_markdown(make_chart())
        assert "Claim Chart — US10123456B2, Claim 1" in md
        # bold bracket label prefixes the verbatim limitation text
        assert "**1[a]** a frame supporting the housing" in md
        assert "**1[b]** a motor coupled to the frame" in md
        assert "US111 (Prior Widget)" in md
        assert "US222" in md
        assert "**Disclosed**" in md
        assert "**Not disclosed**" in md
        assert "**Partial**" in md
        assert "a frame 12 supports the housing 14" in md
        assert "(col. 2, ll. 10-12)" in md

    def test_coverage_section(self):
        md = claim_chart_markdown(make_chart())
        assert "- US111: 50% of limitations disclosed" in md
        assert "- US222: 50% of limitations disclosed" in md
        assert "Combined (any reference): 100%" in md

    def test_html_rendering(self):
        html = claim_chart_html(make_chart())
        assert "<table" in html
        assert "<b>1[a]</b> a frame supporting the housing" in html
        assert "background-color:#c6efce" in html  # disclosed = green
        assert "col. 2, ll. 10-12" in html


class TestClaimChartDocx:
    def test_citation_on_own_paragraph_once(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document

        from patentkit.formatting import claim_chart_docx

        out = tmp_path / "chart.docx"
        claim_chart_docx(make_chart(), str(out))

        document = Document(str(out))
        table = document.tables[0]
        # row 1 (first limitation), column 1 (US111): the cell with a citation
        cell = table.rows[1].cells[1]
        citation = "col. 2, ll. 10-12"
        cell_text = "\n".join(p.text for p in cell.paragraphs)
        assert cell_text.count(citation) == 1
        # the citation sits on its OWN paragraph, not appended to the quote
        cite_pars = [p for p in cell.paragraphs if citation in p.text]
        assert len(cite_pars) == 1
        cite_par = cite_pars[0]
        assert cite_par.text == f"[{citation}]"
        quote_par = next(p for p in cell.paragraphs if "a frame 12 supports" in p.text)
        assert citation not in quote_par.text
        # typography: citation 8pt non-italic; quote 9pt italic with spacing
        from docx.shared import Pt

        assert cite_par.runs[0].font.size == Pt(8)
        assert not cite_par.runs[0].italic
        assert cite_par.paragraph_format.space_after == Pt(6)
        assert quote_par.runs[0].italic
        assert quote_par.runs[0].font.size == Pt(9)
        assert quote_par.paragraph_format.space_before == Pt(4)
        assert quote_par.paragraph_format.space_after == Pt(4)


class TestInvalidityReport:
    def test_report_sections(self):
        patent = Patent(
            patent_number=PatentNumber.parse("US10123456B2"), title="Widget assembly"
        )
        results = [
            SearchResult(
                patent_number=PatentNumber.parse("US111"),
                score=1.25,
                passages=[Passage(text="a highlighted prior-art passage")],
            )
        ]
        md = invalidity_report_markdown(
            patent, results, charts=[make_chart()], params={"keywords": "frame, motor"}
        )
        assert "# Invalidity Search Report — US10123456B2" in md
        assert "## Executive Summary" in md
        assert "## Methodology" in md
        assert "| keywords | frame, motor |" in md
        assert "US111" in md
        assert "(score 1.25)" in md
        assert "> a highlighted prior-art passage" in md
        assert "## Appendix: Claim Charts" in md


class TestFtoReport:
    def test_report_contents(self):
        fto = FtoReportData(
            product_description="A smart mug with a heating coil.",
            findings=[
                FtoFinding(
                    patent_number="US333", claim_number=1, risk="literal",
                    confidence=3, assumptions="Coil is resistive.", argument="Every element met.",
                ),
                FtoFinding(
                    patent_number="US444", claim_number=1, risk="none",
                    confidence=2, assumptions="No vacuum layer.",
                ),
            ],
            searched_at="2026-06-09T12:00:00Z",
            query_params={"jurisdiction": "US"},
        )
        assert fto.risk_summary() == {"literal": 1, "doe": 0, "none": 1}
        md = fto_report_markdown(fto)
        assert "# Freedom-to-Operate Report" in md
        assert "A smart mug with a heating coil." in md
        assert "| Literal infringement | 1 |" in md
        assert "| US333 | 1 | literal | 3/3 | Coil is resistive. |" in md
        assert "Every element met." in md
        assert "| jurisdiction | US |" in md


class TestInfringementReport:
    def test_report_contents(self):
        patent = Patent(patent_number=PatentNumber.parse("US555"), title="Pump")
        findings = [
            InfringementFinding(
                limitation=Limitation(label="1[a]", text="a sealed impeller"),
                status="met",
                evidence=[
                    EvidenceItem(
                        source_url="https://example.com/spec",
                        quote="features a sealed impeller",
                        note="product page",
                    )
                ],
                reasoning="Product page confirms it.",
            ),
            InfringementFinding(
                limitation=Limitation(label="1[b]", text="a titanium shaft"),
                status="not_met",
                reasoning="Shaft is steel.",
            ),
        ]
        md = infringement_report_markdown(patent, "AquaPump 3000", findings)
        assert "US555 vs. AquaPump 3000" in md
        assert "1/2 limitation(s) shown met" in md
        assert "| a sealed impeller | met |" in md
        assert "features a sealed impeller" in md
        assert "not met" in md
        assert "Shaft is steel." in md
