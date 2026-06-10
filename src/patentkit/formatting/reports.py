"""Markdown + DOCX report generation for invalidity, FTO, and infringement
workflows.

The markdown renderers are dependency-free; ``*_docx`` wrappers convert the
markdown through a small shared markdown->docx helper (headings, tables,
bullets, quotes, paragraphs only) using python-docx lazily.
"""

from __future__ import annotations

import logging
import re
from typing import Optional

from patentkit.analysis.fto import FtoReportData
from patentkit.analysis.infringement import InfringementFinding
from patentkit.analysis.invalidity import ClaimChart
from patentkit.formatting.claim_chart import claim_chart_markdown
from patentkit.models import Patent
from patentkit.search.base import SearchResult

logger = logging.getLogger(__name__)

__all__ = [
    "invalidity_report_markdown",
    "fto_report_markdown",
    "infringement_report_markdown",
    "invalidity_report_docx",
    "fto_report_docx",
    "infringement_report_docx",
]


def _md_cell(text: str) -> str:
    return text.replace("|", "\\|").replace("\n", " ").strip()


def _params_table(params: dict) -> list[str]:
    lines = ["| Parameter | Value |", "| --- | --- |"]
    for key, value in params.items():
        lines.append(f"| {_md_cell(str(key))} | {_md_cell(str(value))} |")
    return lines


# ---------------------------------------------------------------------------
# Invalidity
# ---------------------------------------------------------------------------


def invalidity_report_markdown(
    query_patent: Patent,
    results: list[SearchResult],
    charts: Optional[list[ClaimChart]] = None,
    params: Optional[dict] = None,
) -> str:
    """Render an invalidity (prior-art) search report as markdown.

    Sections: executive summary, methodology/query parameters, ranked results
    with highlighted passages, and a claim-chart appendix.
    """
    number = str(query_patent.patent_number)
    lines = [f"# Invalidity Search Report — {number}", ""]
    if query_patent.title:
        lines += [f"**Title:** {query_patent.title}", ""]

    lines += ["## Executive Summary", ""]
    lines.append(
        f"This report covers a prior-art search against {number}, returning "
        f"{len(results)} candidate reference(s)."
    )
    if charts:
        for chart in charts:
            lines.append(
                f"Claim {chart.claim_number}: best single-reference coverage "
                f"{max(chart.coverage_summary().values(), default=0.0):.0%}; combined "
                f"coverage {chart.combined_coverage():.0%}."
            )
    lines.append("")

    lines += ["## Methodology", ""]
    if params:
        lines += _params_table(params)
    else:
        lines.append("_No query parameters recorded._")
    lines.append("")

    lines += ["## Ranked Results", ""]
    for i, result in enumerate(results, start=1):
        title = f" — {result.title}" if result.title else ""
        lines.append(f"### {i}. {result.patent_number}{title} (score {result.score:.2f})")
        if result.explanation:
            lines.append(f"_{result.explanation}_")
        for passage in result.passages:
            lines.append(f"> {passage.text}")
        lines.append("")

    if charts:
        lines += ["## Appendix: Claim Charts", ""]
        for chart in charts:
            lines += [claim_chart_markdown(chart), ""]
    return "\n".join(lines).strip() + "\n"


# ---------------------------------------------------------------------------
# FTO
# ---------------------------------------------------------------------------


def fto_report_markdown(fto: FtoReportData) -> str:
    """Render a freedom-to-operate report as markdown."""
    lines = ["# Freedom-to-Operate Report", ""]
    if fto.searched_at:
        lines += [f"_Search performed: {fto.searched_at}_", ""]
    lines += ["## Product Description", "", fto.product_description, ""]

    summary = fto.risk_summary()
    lines += [
        "## Risk Summary",
        "",
        "| Risk level | Findings |",
        "| --- | --- |",
        f"| Literal infringement | {summary['literal']} |",
        f"| Doctrine of equivalents | {summary['doe']} |",
        f"| No infringement | {summary['none']} |",
        "",
    ]

    if fto.query_params:
        lines += ["## Search Parameters", ""] + _params_table(fto.query_params) + [""]

    lines += ["## Findings", ""]
    if not fto.findings:
        lines.append("_No findings._")
    else:
        lines += [
            "| Patent | Claim | Risk | Confidence | Key assumptions |",
            "| --- | --- | --- | --- | --- |",
        ]
        for f in fto.findings:
            lines.append(
                f"| {f.patent_number} | {f.claim_number} | {f.risk} | {f.confidence}/3 "
                f"| {_md_cell(f.assumptions)} |"
            )
        lines.append("")
        argued = [f for f in fto.findings if f.argument]
        if argued:
            lines += ["### Detailed Arguments", ""]
            for f in argued:
                lines += [f"**{f.patent_number}, claim {f.claim_number} ({f.risk}):**", f.argument, ""]
    return "\n".join(lines).strip() + "\n"


# ---------------------------------------------------------------------------
# Infringement
# ---------------------------------------------------------------------------


def infringement_report_markdown(
    patent: Patent, product_name: str, findings: list[InfringementFinding]
) -> str:
    """Render a per-limitation infringement evidence report as markdown."""
    lines = [f"# Infringement Analysis — {patent.patent_number} vs. {product_name}", ""]
    if patent.title:
        lines += [f"**Patent title:** {patent.title}", ""]

    met = sum(1 for f in findings if f.status == "met")
    likely = sum(1 for f in findings if f.status in ("met", "likely"))
    lines += [
        "## Summary",
        "",
        f"{met}/{len(findings)} limitation(s) shown met by the evidence; "
        f"{likely}/{len(findings)} met or likely met.",
        "",
        "## Limitation-by-Limitation Findings",
        "",
        "| Limitation | Status | Evidence | Reasoning |",
        "| --- | --- | --- | --- |",
    ]
    for f in findings:
        evidence = "<br>".join(
            f"“{_md_cell(e.quote)}” ({_md_cell(e.source_url)})" for e in f.evidence if e.quote
        ) or "—"
        lines.append(
            f"| {_md_cell(f.limitation.text)} | {f.status.replace('_', ' ')} "
            f"| {evidence} | {_md_cell(f.reasoning)} |"
        )
    return "\n".join(lines).strip() + "\n"


# ---------------------------------------------------------------------------
# Shared markdown -> DOCX helper and wrappers
# ---------------------------------------------------------------------------

_EMPHASIS_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_|`(.+?)`")


def _strip_emphasis(text: str) -> str:
    return _EMPHASIS_RE.sub(lambda m: next(g for g in m.groups() if g is not None), text)


def _markdown_to_docx(markdown: str, out_path: str) -> None:
    """Convert simple report markdown (headings/tables/bullets/quotes/paragraphs)
    to a DOCX file. Requires the ``docx`` extra."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX reports; install with: pip install 'patentkit[docx]'"
        ) from exc

    document = Document()
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            document.add_heading(_strip_emphasis(heading.group(2)), level=len(heading.group(1)))
            i += 1
            continue
        if line.lstrip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [
                [_strip_emphasis(c.strip()).replace("<br>", "\n") for c in row.strip("|").split("|")]
                for row in table_lines
                if not re.fullmatch(r"\|?[\s:|-]+\|?", row)
            ]
            if rows:
                table = document.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        table.cell(r, c).text = cell_text
                    if r == 0:
                        for cell in table.rows[0].cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
            continue
        if line.lstrip().startswith("- "):
            document.add_paragraph(_strip_emphasis(line.lstrip()[2:]), style="List Bullet")
            i += 1
            continue
        if line.lstrip().startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.add_run(_strip_emphasis(line.lstrip()[2:])).italic = True
            i += 1
            continue
        document.add_paragraph(_strip_emphasis(line))
        i += 1
    document.save(out_path)
    logger.info("Wrote report to %s", out_path)


def invalidity_report_docx(
    query_patent: Patent,
    results: list[SearchResult],
    out_path: str,
    charts: Optional[list[ClaimChart]] = None,
    params: Optional[dict] = None,
) -> None:
    """DOCX wrapper around :func:`invalidity_report_markdown` (``docx`` extra)."""
    _markdown_to_docx(invalidity_report_markdown(query_patent, results, charts, params), out_path)


def fto_report_docx(fto: FtoReportData, out_path: str) -> None:
    """DOCX wrapper around :func:`fto_report_markdown` (``docx`` extra)."""
    _markdown_to_docx(fto_report_markdown(fto), out_path)


def infringement_report_docx(
    patent: Patent, product_name: str, findings: list[InfringementFinding], out_path: str
) -> None:
    """DOCX wrapper around :func:`infringement_report_markdown` (``docx`` extra)."""
    _markdown_to_docx(infringement_report_markdown(patent, product_name, findings), out_path)
