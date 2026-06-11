"""Render :class:`~patentkit.analysis.invalidity.ClaimChart` objects as
markdown, HTML, and DOCX claim charts.

Markdown and HTML renderers are dependency-free; the DOCX renderer imports
python-docx lazily (``docx`` extra) and color-codes cells by disclosure
status when requested.
"""

from __future__ import annotations

import html as _html
import logging
from typing import Optional, Sequence

from patentkit.analysis.invalidity import ClaimChart, DisclosureFinding, ReferenceChart
from patentkit.models import AtomicLimitation

logger = logging.getLogger(__name__)

__all__ = [
    "claim_chart_markdown",
    "claim_chart_html",
    "claim_chart_docx",
    "filter_chart",
    "limitation_chart_markdown",
]

_STATUS_LABEL = {
    "disclosed": "Disclosed",
    "partial": "Partial",
    "not_disclosed": "Not disclosed",
}

#: cell shading per status (hex fill, no '#')
_STATUS_FILL = {
    "disclosed": "C6EFCE",   # green
    "partial": "FFEB9C",     # yellow
    "not_disclosed": "FFC7CE",  # red
}

_STATUS_CSS = {
    "disclosed": "#c6efce",
    "partial": "#ffeb9c",
    "not_disclosed": "#ffc7ce",
}


def _finding_for(
    reference: ReferenceChart, limitation: AtomicLimitation, index: int
) -> Optional[DisclosureFinding]:
    """Find the reference's finding for ``limitation`` (by position, then text)."""
    if index < len(reference.findings) and reference.findings[index].limitation.text == limitation.text:
        return reference.findings[index]
    return next((f for f in reference.findings if f.limitation.text == limitation.text), None)


def _ref_header(reference: ReferenceChart) -> str:
    if reference.reference_title:
        return f"{reference.reference_number} ({reference.reference_title})"
    return reference.reference_number


def _md_escape(text: str) -> str:
    return text.replace("|", "\\|").replace("\n", "<br>")


def _finding_cell_markdown(finding: Optional[DisclosureFinding]) -> str:
    if finding is None:
        return "—"
    parts = [f"**{_STATUS_LABEL[finding.status]}**"]
    if finding.reasoning:
        parts.append(_md_escape(finding.reasoning))
    for i, quote in enumerate(finding.quotes):
        # the citation was located from the first quote only (see assess_reference)
        cite = f" ({finding.citation})" if finding.citation and i == 0 else ""
        parts.append(f"“{_md_escape(quote)}”{cite}")
    return "<br>".join(parts)


def claim_chart_markdown(chart: ClaimChart) -> str:
    """Render a claim chart as a dependency-free markdown table."""
    lines = [f"## Claim Chart — {chart.query_patent}, Claim {chart.claim_number}", ""]
    if chart.interpreted_claim:
        lines += ["**Interpreted claim:** " + _md_escape(chart.interpreted_claim), ""]

    headers = ["Limitation"] + [_md_escape(_ref_header(ref)) for ref in chart.references]
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("|" + "|".join([" --- "] * len(headers)) + "|")
    for i, limitation in enumerate(chart.limitations):
        cells = [_md_escape(limitation.text)]
        for ref in chart.references:
            cells.append(_finding_cell_markdown(_finding_for(ref, limitation, i)))
        lines.append("| " + " | ".join(cells) + " |")

    lines += ["", "### Coverage"]
    for number, fraction in chart.coverage_summary().items():
        lines.append(f"- {number}: {fraction:.0%} of limitations disclosed")
    lines.append(f"- Combined (any reference): {chart.combined_coverage():.0%}")
    return "\n".join(lines)


def filter_chart(chart: ClaimChart, limitation_substrings: Sequence[str]) -> ClaimChart:
    """Copy of ``chart`` restricted to limitations matching any substring.

    Matching is case-insensitive substring containment against each
    limitation's text. The returned chart's coverage figures reflect only the
    retained rows. Limitations matching nothing are dropped; an empty filter
    list returns the chart unchanged.
    """
    wanted = [str(s).strip().lower() for s in limitation_substrings if str(s).strip()]
    if not wanted:
        return chart
    matched = [lim for lim in chart.limitations
               if any(w in lim.text.lower() for w in wanted)]
    texts = {lim.text for lim in matched}
    references = [
        ReferenceChart(
            reference_number=ref.reference_number,
            reference_title=ref.reference_title,
            findings=[f for f in ref.findings if f.limitation.text in texts],
        )
        for ref in chart.references
    ]
    return ClaimChart(
        query_patent=chart.query_patent,
        claim_number=chart.claim_number,
        interpreted_claim=chart.interpreted_claim,
        limitations=matched,
        references=references,
    )


def limitation_chart_markdown(chart: ClaimChart, limitation: AtomicLimitation) -> str:
    """Render one limitation's disclosure across references: one row per
    reference with status, reasoning, and quotes/citation."""
    lines = [
        f"## Limitation Chart — {chart.query_patent}, Claim {chart.claim_number}",
        "",
        f"**Limitation:** {_md_escape(limitation.text)}",
        "",
        "| Reference | Status | Reasoning | Quotes |",
        "| --- | --- | --- | --- |",
    ]
    for ref in chart.references:
        finding = next((f for f in ref.findings if f.limitation.text == limitation.text), None)
        if finding is None:
            lines.append(f"| {_md_escape(_ref_header(ref))} | — | — | — |")
            continue
        quotes = []
        for i, quote in enumerate(finding.quotes):
            cite = f" ({finding.citation})" if finding.citation and i == 0 else ""
            quotes.append(f"“{_md_escape(quote)}”{cite}")
        lines.append(
            f"| {_md_escape(_ref_header(ref))} "
            f"| **{_STATUS_LABEL[finding.status]}** "
            f"| {_md_escape(finding.reasoning) or '—'} "
            f"| {'<br>'.join(quotes) or '—'} |"
        )
    return "\n".join(lines)


def claim_chart_html(chart: ClaimChart) -> str:
    """Render a claim chart as a standalone HTML table (status color-coded)."""
    esc = _html.escape
    rows = []
    header_cells = "".join(
        f"<th>{esc(h)}</th>" for h in ["Limitation"] + [_ref_header(r) for r in chart.references]
    )
    rows.append(f"<tr>{header_cells}</tr>")
    for i, limitation in enumerate(chart.limitations):
        cells = [f"<td>{esc(limitation.text)}</td>"]
        for ref in chart.references:
            finding = _finding_for(ref, limitation, i)
            if finding is None:
                cells.append("<td>—</td>")
                continue
            quote_html = "".join(
                f"<blockquote>{esc(q)}"
                f"{(' <i>(' + esc(finding.citation) + ')</i>') if finding.citation and qi == 0 else ''}"
                "</blockquote>"
                for qi, q in enumerate(finding.quotes)
            )
            cells.append(
                f'<td style="background-color:{_STATUS_CSS[finding.status]}">'
                f"<b>{_STATUS_LABEL[finding.status]}</b>"
                f"<p>{esc(finding.reasoning)}</p>{quote_html}</td>"
            )
        rows.append("<tr>" + "".join(cells) + "</tr>")
    table = "\n".join(rows)
    return (
        f"<h2>Claim Chart — {esc(chart.query_patent)}, Claim {chart.claim_number}</h2>\n"
        f'<table border="1" cellspacing="0" cellpadding="6">\n{table}\n</table>'
    )


def claim_chart_docx(chart: ClaimChart, out_path: str, color_coding: bool = True) -> None:
    """Write a claim chart as a DOCX table (requires the ``docx`` extra).

    One row per limitation; one column per reference. With ``color_coding``,
    cells are shaded green/yellow/red by disclosure status. Quotes are
    rendered with their citations (e.g. "col. 3, ll. 45-52") when present.
    """
    try:
        from docx import Document
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.shared import Pt
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX claim charts; install with: pip install 'patentkit[docx]'"
        ) from exc

    def shade(cell, fill: str) -> None:
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
        )

    document = Document()
    document.add_heading(f"Claim Chart — {chart.query_patent}, Claim {chart.claim_number}", level=1)
    if chart.interpreted_claim:
        document.add_paragraph(chart.interpreted_claim)

    table = document.add_table(rows=1, cols=1 + len(chart.references))
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Limitation"
    for j, ref in enumerate(chart.references):
        header[j + 1].text = _ref_header(ref)
    for cell in header:
        shade(cell, "DCDCDC")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for i, limitation in enumerate(chart.limitations):
        row = table.add_row().cells
        row[0].text = limitation.text
        for j, ref in enumerate(chart.references):
            cell = row[j + 1]
            finding = _finding_for(ref, limitation, i)
            if finding is None:
                cell.text = "—"
                continue
            status_par = cell.paragraphs[0]
            status_par.add_run(_STATUS_LABEL[finding.status]).bold = True
            status_par.paragraph_format.space_after = Pt(4)
            if finding.reasoning:
                reasoning_par = cell.add_paragraph(finding.reasoning)
                reasoning_par.paragraph_format.space_after = Pt(6)
            for qi, quote in enumerate(finding.quotes):
                par = cell.add_paragraph()
                par.paragraph_format.space_before = Pt(4)
                par.paragraph_format.space_after = Pt(4)
                run = par.add_run(f"“{quote}”")
                run.italic = True
                run.font.size = Pt(9)
                if finding.citation and qi == 0:
                    # the citation was located from the first quote only
                    # (see assess_reference); give it its own paragraph
                    cite_par = cell.add_paragraph()
                    cite_par.paragraph_format.space_after = Pt(6)
                    cite_run = cite_par.add_run(f"[{finding.citation}]")
                    cite_run.italic = False
                    cite_run.font.size = Pt(8)
            if color_coding:
                shade(cell, _STATUS_FILL[finding.status])

    document.add_paragraph()
    summary = document.add_paragraph("Coverage: ")
    parts = [
        f"{number}: {fraction:.0%}" for number, fraction in chart.coverage_summary().items()
    ]
    parts.append(f"combined: {chart.combined_coverage():.0%}")
    summary.add_run("; ".join(parts))
    document.save(out_path)
    logger.info("Wrote claim chart to %s", out_path)
