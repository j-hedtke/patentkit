"""Generic document text extraction with lazy optional backends.

``extract_text`` dispatches on file type: PDF via pymupdf (``pdf`` extra),
DOCX via python-docx (``docx`` extra), HTML via beautifulsoup4 (``scrape``
extra) with a zero-dependency stdlib fallback, and plain text/markdown read
directly.
"""

from __future__ import annotations

import logging
from html.parser import HTMLParser
from pathlib import Path
from typing import Optional, Union

logger = logging.getLogger(__name__)

__all__ = ["extract_text", "html_to_text"]

_PDF_MAGIC = b"%PDF"
_ZIP_MAGIC = b"PK\x03\x04"  # docx is a zip container


def extract_text(path_or_bytes: Union[str, Path, bytes], kind: Optional[str] = None) -> str:
    """Extract plain text from a document.

    Args:
        path_or_bytes: a filesystem path (str/Path) or raw document bytes.
        kind: explicit type override ("pdf", "docx", "html", "txt", "md");
            inferred from the file suffix (or magic bytes for raw bytes)
            when omitted.

    Raises:
        ImportError: if the required optional backend is not installed
            (the message names the pip extra to install).
        ValueError: if the document type cannot be determined.
    """
    if isinstance(path_or_bytes, bytes):
        data = path_or_bytes
        kind = kind or _sniff_kind(data)
    else:
        path = Path(path_or_bytes)
        kind = kind or path.suffix.lstrip(".").lower()
        if kind in ("txt", "md", "text", "markdown", ""):
            return path.read_text(encoding="utf-8", errors="replace")
        data = path.read_bytes()

    kind = (kind or "").lower().lstrip(".")
    if kind == "pdf":
        return _pdf_to_text(data)
    if kind == "docx":
        return _docx_to_text(data)
    if kind in ("html", "htm"):
        return html_to_text(data.decode("utf-8", errors="replace"))
    if kind in ("txt", "md", "text", "markdown"):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Cannot determine document type (kind={kind!r}); pass kind= explicitly")


def _sniff_kind(data: bytes) -> str:
    if data[:4] == _PDF_MAGIC:
        return "pdf"
    if data[:4] == _ZIP_MAGIC:
        return "docx"
    head = data[:512].lstrip().lower()
    if head.startswith(b"<!doctype html") or head.startswith(b"<html"):
        return "html"
    return "txt"


def _pdf_to_text(data: bytes) -> str:
    try:
        import fitz  # type: ignore[import-not-found]
    except ImportError as exc:
        raise ImportError(
            "pymupdf is required for PDF extraction; install with: pip install 'patentkit[pdf]'"
        ) from exc
    with fitz.open(stream=data, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def _docx_to_text(data: bytes) -> str:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX extraction; install with: pip install 'patentkit[docx]'"
        ) from exc
    from io import BytesIO

    document = docx.Document(BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(p for p in parts if p)


_BLOCK_TAGS = {
    "p", "div", "br", "li", "tr", "table", "section", "article", "header",
    "footer", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "blockquote", "pre",
}
_SKIP_TAGS = {"script", "style", "head", "noscript", "template"}


class _TextExtractor(HTMLParser):
    """Stdlib HTML-to-text converter: skips script/style, breaks on block tags."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._chunks: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in _SKIP_TAGS:
            self._skip_depth += 1
        elif tag in _BLOCK_TAGS:
            self._chunks.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in _SKIP_TAGS and self._skip_depth > 0:
            self._skip_depth -= 1
        elif tag in _BLOCK_TAGS:
            self._chunks.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0 and data:
            self._chunks.append(data)

    def text(self) -> str:
        raw = "".join(self._chunks)
        lines = [" ".join(line.split()) for line in raw.splitlines()]
        out: list[str] = []
        for line in lines:
            if line:
                out.append(line)
            elif out and out[-1]:
                out.append("")
        return "\n".join(out).strip()


def html_to_text(html: str) -> str:
    """Convert HTML to plain text.

    Uses beautifulsoup4 when available (``scrape`` extra) and falls back to a
    stdlib :class:`html.parser.HTMLParser` implementation, so it always works
    with zero optional dependencies.
    """
    try:
        from bs4 import BeautifulSoup  # type: ignore[import-not-found]
    except ImportError:
        parser = _TextExtractor()
        parser.feed(html)
        return parser.text()
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "template"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
